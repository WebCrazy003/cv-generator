import tempfile
from pathlib import Path

from docx import Document

import app

TEMPLATE = Path(__file__).resolve().parent.parent / "cv_template.docx"

SAMPLE_DATA = {
    "personNameOnCV": "Jane Doe",
    "personLocation": "Berlin, Germany",
    "personEmail": "jane.doe@example.com",
    "personUniversity": "TU Berlin",
    "personDegree": "BSc Computer Science",
    "companyNameApplyJob": "Acme Corp",
    "summary": "Senior engineer with a focus on ML platforms.",
    "experience": [
        {
            "companyName": "SoftXPro",
            "jobTitle": "Senior ML Engineer",
            "startDate": "11/2023",
            "endDate": "07/2026",
            "content": ["Built LLM applications", "Shipped inference services"],
        },
        {
            "companyName": "STS Software",
            "jobTitle": "Backend Engineer",
            "startDate": "01/2020",
            "endDate": "10/2023",
            "content": ["Distributed backend services"],
        },
    ],
    "skills": [
        {"categoryName": "Programming Languages", "skillItems": ["Python", "SQL"]},
        {"categoryName": "MLOps", "skillItems": ["Docker", "Kubernetes"]},
    ],
}


def _generate():
    tmp_dir = tempfile.mkdtemp()
    output = Path(tmp_dir) / "out.docx"
    app.build_docx(TEMPLATE, SAMPLE_DATA, output)
    text = "\n".join(p.text for p in Document(str(output)).paragraphs)
    return output, text


def test_output_has_no_leftover_tokens():
    _, text = _generate()
    assert "{{" not in text and "}}" not in text


def test_output_contains_injected_content():
    _, text = _generate()
    for expected in ["Jane Doe", "SoftXPro", "Backend Engineer",
                     "Programming Languages: Python, SQL"]:
        assert expected in text, expected


def test_output_contains_contact_and_education_fields():
    _, text = _generate()
    # Location and email are written together on the contact line.
    assert "Berlin, Germany • jane.doe@example.com" in text
    # University and degree are written together on the education line.
    assert "TU Berlin | BSc Computer Science" in text


def test_output_drops_original_sample_content():
    _, text = _generate()
    for sample in ["Robson", "SDLC Corp", "Pitágoras"]:
        assert sample not in text, sample


def test_body_font_normalized_for_pdf():
    output, _ = _generate()
    from docx.oxml.ns import qn

    fonts = set()
    for rfonts in Document(str(output)).element.iter(qn("w:rFonts")):
        for attr in ("w:ascii", "w:hAnsi"):
            value = rfonts.get(qn(attr))
            if value:
                fonts.add(value)
    # "Avenir Book" does not resolve in LibreOffice (falls back to serif); it is
    # remapped to "Avenir" so the PDF renders the intended sans-serif.
    assert "Avenir Book" not in fonts
    assert "Avenir" in fonts


def test_section_styles_preserved():
    output, _ = _generate()
    from docx.oxml.ns import qn

    header_styles = []
    for paragraph in Document(str(output)).paragraphs:
        ppr = paragraph._p.find(qn("w:pPr"))
        if ppr is None:
            continue
        pstyle = ppr.find(qn("w:pStyle"))
        if pstyle is not None and pstyle.get(qn("w:val")) == "2":
            header_styles.append(paragraph.text)
    assert header_styles == ["Experience", "Skills", "Education"]


def test_name_paragraph_centered_in_box():
    output, _ = _generate()
    from docx.oxml.ns import qn

    ppr = Document(str(output)).paragraphs[0]._p.find(qn("w:pPr"))
    spacing = ppr.find(qn("w:spacing"))
    # Exact line height makes LibreOffice vertically center the name in its box.
    assert spacing is not None
    assert spacing.get(qn("w:lineRule")) == "exact"
    assert spacing.get(qn("w:before")) == str(app.NAME_TOP_PADDING_TWIPS)


def test_handles_single_experience_entry():
    data = dict(SAMPLE_DATA)
    data["experience"] = [SAMPLE_DATA["experience"][0]]
    tmp_dir = tempfile.mkdtemp()
    output = Path(tmp_dir) / "single.docx"
    app.build_docx(TEMPLATE, data, output)
    text = "\n".join(p.text for p in Document(str(output)).paragraphs)
    assert "STS Software" not in text
    assert "SoftXPro" in text
