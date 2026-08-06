import copy
import json
import mimetypes
import os
import re
import shutil
import subprocess
from datetime import datetime
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT_DIR = Path(__file__).resolve().parent
FRONTEND_DIR = ROOT_DIR / "frontend"
SETTINGS_FILE = ROOT_DIR / "settings.json"
TEMPLATE_FILE = ROOT_DIR / "cv_template.docx"

DEFAULT_SETTINGS = {"outputDirectory": ""}


def load_settings():
    if SETTINGS_FILE.exists():
        try:
            return json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return dict(DEFAULT_SETTINGS)
    return dict(DEFAULT_SETTINGS)


def save_settings(settings):
    SETTINGS_FILE.write_text(json.dumps(settings, indent=2), encoding="utf-8")
    return settings


def sanitize_name(value):
    value = re.sub(r"[^A-Za-z0-9._ -]+", "_", str(value).strip())
    value = re.sub(r"\s+", " ", value).strip()
    return value or "cv"


# ---------------------------------------------------------------------------
# DOCX generation
#
# The template (cv_template.docx) is a fully-formatted sample CV. Rather than
# string-replacing tokens (which fails because tokens are split across runs and
# are followed by hardcoded content), we treat the template as a *style carrier*:
# we harvest one "prototype" paragraph per element type, wipe the body, then
# rebuild the document from the JSON data by cloning those prototypes. Cloning
# preserves the exact fonts, colors, list numbering and paragraph styles of the
# original template.
# ---------------------------------------------------------------------------

MARKER_RE = re.compile(r"\{\{[^}]*\}\}|\[\[[^\]]*\]\]")

# LibreOffice (used for PDF conversion) cannot resolve the macOS font family
# name "Avenir Book" and silently falls back to a serif (Liberation Serif),
# which does not match the template. The plain "Avenir" family name resolves to
# the same typeface (Avenir Roman), so remap it before conversion. Only the PDF
# is kept, so rewriting the font name in the intermediate DOCX is harmless.
FONT_REMAP = {"Avenir Book": "Avenir"}

# LibreOffice floats the large first line (the name) ~8pt above the page's top
# margin, so the PDF looks top-cramped compared to Word, which clamps the first
# line to the margin. Add a little space before the name so the PDF starts at
# the top margin like the DOCX does. Tuned to the template's name size (35pt).
NAME_TOP_PADDING_TWIPS = 200

# The name sits in a shaded (highlight) box whose height is the line box. With
# the template's default "auto" line rule LibreOffice top-aligns the name, so
# the box has no space above the text and a large empty gap below (the font's
# descent region, unused because the name has no descenders). Switching to an
# "exact" line height near the natural line height makes LibreOffice vertically
# center the text in the box instead. Tuned to the template's name size (35pt).
NAME_LINE_HEIGHT_TWIPS = 800


def normalize_fonts(document, mapping=FONT_REMAP):
    for rfonts in document.element.iter(qn("w:rFonts")):
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            value = rfonts.get(qn(attr))
            if value in mapping:
                rfonts.set(qn(attr), mapping[value])


def _pstyle_id(paragraph):
    ppr = paragraph._p.find(qn("w:pPr"))
    if ppr is None:
        return None
    pstyle = ppr.find(qn("w:pStyle"))
    return pstyle.get(qn("w:val")) if pstyle is not None else None


def _is_list(paragraph):
    ppr = paragraph._p.find(qn("w:pPr"))
    return ppr is not None and ppr.find(qn("w:numPr")) is not None


def _strip_markers(text):
    return MARKER_RE.sub("", text)


def harvest_prototypes(document):
    """Locate one representative paragraph element for each CV element type."""
    protos = {}
    skills_idx = education_idx = None
    paragraphs = document.paragraphs

    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text
        pstyle = _pstyle_id(paragraph)
        if "{{Skills}}" in text:
            skills_idx = index
        if "{{Education}}" in text:
            education_idx = index
        if "name" not in protos and index == 0:
            protos["name"] = paragraph._p
        if "contact" not in protos and pstyle == "91":
            protos["contact"] = paragraph._p
        if "summary" not in protos and "{{summary}}" in text:
            protos["summary"] = paragraph._p
        if "header" not in protos and pstyle == "2":
            protos["header"] = paragraph._p
        if "date" not in protos and pstyle == "3":
            protos["date"] = paragraph._p
        if "title" not in protos and pstyle == "4" and "{{company" in text:
            protos["title"] = paragraph._p
        if "bullet" not in protos and _is_list(paragraph):
            protos["bullet"] = paragraph._p
        if "spacer" not in protos and not paragraph.runs and pstyle is None:
            protos["spacer"] = paragraph._p

    # Skill lines and education lines share pStyle/list traits with bullets/titles
    # but carry different inline formatting, so harvest them from their sections.
    if skills_idx is not None:
        for paragraph in paragraphs[skills_idx + 1:]:
            if _is_list(paragraph):
                protos["skill"] = paragraph._p
                break
    if education_idx is not None:
        for paragraph in paragraphs[education_idx + 1:]:
            if _pstyle_id(paragraph) == "4":
                protos["education"] = paragraph._p
                break

    # Deep-copy so later body wiping cannot disturb the prototypes.
    return {key: copy.deepcopy(element) for key, element in protos.items()}


def _clone_paragraph(proto_element, text):
    """Clone a prototype paragraph, collapsing it to a single run with `text`."""
    new_p = copy.deepcopy(proto_element)
    runs = new_p.findall(qn("w:r"))
    for extra in runs[1:]:
        new_p.remove(extra)
    if runs:
        first = runs[0]
        for child in list(first):
            if child.tag != qn("w:rPr"):
                first.remove(child)
        text_el = OxmlElement("w:t")
        text_el.set(qn("xml:space"), "preserve")
        text_el.text = text
        first.append(text_el)
    return new_p


def build_docx(template_path, data, output_path):
    template_path = Path(template_path).expanduser()
    output_path = Path(output_path).expanduser()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if not template_path.exists():
        raise FileNotFoundError(f"Template path does not exist: {template_path}")

    document = Document(str(template_path))
    protos = harvest_prototypes(document)

    body = document.element.body
    sect_pr = body.find(qn("w:sectPr"))

    # Wipe existing body content (keep section properties).
    for child in list(body):
        if child.tag in (qn("w:p"), qn("w:tbl")):
            body.remove(child)

    def emit(proto_key, text):
        proto = protos.get(proto_key)
        if proto is None:
            return
        element = _clone_paragraph(proto, text)
        if sect_pr is not None:
            sect_pr.addprevious(element)
        else:
            body.append(element)

    def emit_spacer():
        proto = protos.get("spacer")
        if proto is None:
            return
        element = copy.deepcopy(proto)
        if sect_pr is not None:
            sect_pr.addprevious(element)
        else:
            body.append(element)

    name = str(data.get("personNameOnCV") or "").strip()
    location = str(data.get("personLocation") or "").strip()
    email = str(data.get("personEmail") or "").strip()
    contact = " • ".join(part for part in (location, email) if part)
    summary = str(data.get("summary") or "").strip()
    experience = data.get("experience") or []
    skills = data.get("skills") or []
    university = str(data.get("personUniversity") or "").strip()
    degree = str(data.get("personDegree") or "").strip()

    if name:
        emit("name", name)
    if contact:
        emit("contact", contact)
    emit_spacer()
    if summary:
        emit("summary", _strip_markers(summary))

    if experience:
        emit("header", "Experience")
        for position, item in enumerate(experience):
            start = str(item.get("startDate") or "").strip()
            end = str(item.get("endDate") or "").strip()
            date_text = f"{start} – {end}".strip(" –") if (start or end) else ""
            job_title = str(item.get("jobTitle") or "").strip()
            company = str(item.get("companyName") or "").strip()
            title_text = " | ".join(part for part in (job_title, company) if part)
            if date_text:
                emit("date", date_text)
            if title_text:
                emit("title", title_text)
            for bullet in item.get("content") or []:
                emit("bullet", str(bullet))
            if position != len(experience) - 1:
                emit_spacer()

    if skills:
        emit("header", "Skills")
        for item in skills:
            category = str(item.get("categoryName") or "").strip()
            entries = item.get("skillItems") or []
            if category and entries:
                emit("skill", f"{category}: {', '.join(entries)}")
            elif category:
                emit("skill", category)

    if university or degree:
        emit("header", "Education")
        line = " | ".join(part for part in (university, degree) if part)
        emit("education", line)

    normalize_fonts(document)
    _fix_name_paragraph(document)
    document.save(str(output_path))


def _fix_name_paragraph(document):
    """Adjust the name (first) paragraph for faithful PDF rendering.

    - space before: nudge the whole block down to the page's top margin.
    - exact line height: vertically center the name inside its shaded box.
    """
    paragraphs = document.paragraphs
    if not paragraphs:
        return
    ppr = paragraphs[0]._p.find(qn("w:pPr"))
    if ppr is None:
        return
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:before"), str(NAME_TOP_PADDING_TWIPS))
    spacing.set(qn("w:line"), str(NAME_LINE_HEIGHT_TWIPS))
    spacing.set(qn("w:lineRule"), "exact")


# ---------------------------------------------------------------------------
# PDF generation
# ---------------------------------------------------------------------------

def find_soffice():
    """Return a path to a LibreOffice binary, or None if unavailable."""
    for candidate in ("soffice", "libreoffice"):
        found = shutil.which(candidate)
        if found:
            return found
    mac_path = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    if mac_path.exists():
        return str(mac_path)
    return None


def convert_to_pdf(docx_path, pdf_path):
    """Convert DOCX to PDF via LibreOffice. Falls back to a simple text PDF.

    Returns a warning string when the fallback was used, otherwise None.
    """
    docx_path = Path(docx_path)
    pdf_path = Path(pdf_path)
    soffice = find_soffice()
    if soffice:
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir",
                 str(pdf_path.parent), str(docx_path)],
                check=True,
                capture_output=True,
                timeout=120,
            )
            produced = pdf_path.parent / (docx_path.stem + ".pdf")
            if produced != pdf_path and produced.exists():
                shutil.move(str(produced), str(pdf_path))
            if pdf_path.exists():
                return None
        except (subprocess.SubprocessError, OSError):
            pass  # fall through to fallback
    return build_pdf_fallback(pdf_path, docx_path)


def build_pdf_fallback(pdf_path, docx_path):
    """Write a minimal, dependency-free text PDF from the DOCX paragraph text."""
    try:
        document = Document(str(docx_path))
        lines = [paragraph.text for paragraph in document.paragraphs]
    except Exception:
        lines = []
    build_simple_pdf(pdf_path, lines)
    return ("PDF was generated in a low-fidelity fallback mode. Install LibreOffice "
            "(brew install --cask libreoffice) for a styled PDF matching the DOCX.")


def build_simple_pdf(path, lines):
    content_stream = "BT\n/F1 11 Tf\n72 760 Td\n"
    for line in lines:
        escaped_line = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        content_stream += "(" + escaped_line + ") Tj\n0 -14 Td\n"
    content_stream += "ET"

    content_bytes = content_stream.encode("latin-1", "replace")
    font_obj = b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
    content_obj = b"4 0 obj\n<< /Length " + str(len(content_bytes)).encode("ascii") + b" >>\nstream\n" + content_bytes + b"\nendstream\nendobj\n"
    page_obj = b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
    pages_obj = b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
    catalog_obj = b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
    objects = [catalog_obj, pages_obj, page_obj, content_obj, font_obj]

    pdf_bytes = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf_bytes))
        pdf_bytes.extend(obj)
    xref_offset = len(pdf_bytes)
    pdf_bytes.extend(f"xref\n0 {len(objects)+1}\n".encode("ascii"))
    pdf_bytes.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf_bytes.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf_bytes.extend(f"trailer\n<< /Size {len(objects)+1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii"))
    Path(path).write_bytes(pdf_bytes)


# ---------------------------------------------------------------------------
# HTTP server
# ---------------------------------------------------------------------------

class CvGeneratorHandler(BaseHTTPRequestHandler):
    def _send_json(self, payload, status=200):
        body = json.dumps(payload).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self):
        parsed = urlparse(self.path)
        if parsed.path == "/api/health":
            self._send_json({"status": "ok"})
            return
        if parsed.path == "/api/settings":
            self._send_json(load_settings())
            return

        rel_path = "index.html" if parsed.path in {"/", "/index.html"} else parsed.path.lstrip("/")
        target_path = (FRONTEND_DIR / rel_path).resolve()
        # Guard against path traversal: the resolved path must stay inside FRONTEND_DIR.
        if FRONTEND_DIR.resolve() not in target_path.parents and target_path != FRONTEND_DIR.resolve():
            self._send_json({"error": "Not found"}, 404)
            return
        if target_path.exists() and target_path.is_file():
            self.send_response(200)
            self.send_header("Content-Type", mimetypes.guess_type(str(target_path))[0] or "application/octet-stream")
            self.send_header("Content-Length", str(target_path.stat().st_size))
            self.end_headers()
            self.wfile.write(target_path.read_bytes())
            return
        self._send_json({"error": "Not found"}, 404)

    def do_POST(self):
        parsed = urlparse(self.path)
        content_length = int(self.headers.get("Content-Length", "0"))
        body = self.rfile.read(content_length).decode("utf-8") if content_length else ""

        if parsed.path == "/api/settings":
            try:
                settings = json.loads(body or "{}")
            except json.JSONDecodeError as exc:
                self._send_json({"error": f"Invalid settings JSON: {exc}"}, 400)
                return
            save_settings(settings)
            self._send_json({"status": "saved", "settings": settings})
            return

        if parsed.path == "/api/generate-cv":
            try:
                payload = json.loads(body or "{}")
            except json.JSONDecodeError as exc:
                self._send_json({"error": f"Invalid request JSON: {exc}"}, 400)
                return

            try:
                json_content = payload.get("jsonContent")
                if not json_content:
                    raise ValueError("The JSON content is empty.")
                data = json.loads(json_content)

                if not TEMPLATE_FILE.exists():
                    raise ValueError(f"Template not found: {TEMPLATE_FILE.name}. "
                                     "Place cv_template.docx next to app.py.")

                output_directory = (payload.get("outputDirectory") or "").strip()
                if not output_directory:
                    raise ValueError("Please provide an output base folder.")
                # Persist the base folder so it is remembered across sessions.
                save_settings({"outputDirectory": output_directory})

                output_directory_path = Path(output_directory).expanduser()
                output_directory_path.mkdir(parents=True, exist_ok=True)
                today = datetime.now()
                output_folder = output_directory_path / f"{today:%y_%m_%d}"
                output_folder.mkdir(parents=True, exist_ok=True)

                company_name = sanitize_name(data.get("companyNameApplyJob") or "cv")
                person_name = sanitize_name(data.get("personNameOnCV") or "cv")
                pdf_path = output_folder / f"{person_name}_{company_name}.pdf"
                docx_path = output_folder / f"{person_name}_{company_name}.docx"

                build_docx(TEMPLATE_FILE, data, docx_path)
                warning = convert_to_pdf(docx_path, pdf_path)

                # Only the PDF is wanted; drop the intermediate DOCX once the PDF exists.
                if pdf_path.exists():
                    try:
                        docx_path.unlink()
                    except OSError:
                        pass

                response = {"status": "success", "pdfPath": str(pdf_path)}
                if warning:
                    response["warning"] = warning
                self._send_json(response)
            except Exception as exc:  # pragma: no cover - surfaced to UI
                self._send_json({"error": str(exc)}, 500)
                return
            return

        self._send_json({"error": "Not found"}, 404)


def main():
    server = ThreadingHTTPServer(("127.0.0.1", 8000), CvGeneratorHandler)
    print("Server running on http://127.0.0.1:8000")
    server.serve_forever()


if __name__ == "__main__":
    main()
